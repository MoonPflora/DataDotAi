import json
import re
from docx import Document

def unparse_word(json_text_or_list, output_path):
    """
    Convert parsed Word or Excel/CSV text (JSON array) back into a Word .docx document.

    - Paragraphs are plain text strings, may contain **bold** and *italic* markdown.
    - Tables are sequences of lines like: Table0,row,col,"value" (Word) 
      or Sheet0,row,col,"value" (Excel/CSV).
      Consecutive table lines form a table.
    - Paragraphs and tables can interleave.
    """

    # Accept list directly or decode JSON string
    if isinstance(json_text_or_list, str):
        try:
            lines = json.loads(json_text_or_list)
        except Exception as e:
            raise ValueError(f"Failed to decode JSON: {e}")
    elif isinstance(json_text_or_list, list):
        lines = json_text_or_list
    else:
        raise TypeError("Input must be a JSON string or list")

    doc = Document()
    i = 0
    n = len(lines)

    # Table patterns: Word and Excel
    table_patterns = [
        re.compile(r'Table\d+,\d+,\d+,".*?"'),
        re.compile(r'Sheet\d+,\d+,\d+,".*?"')
    ]

    def is_table_line(line):
        return any(p.match(line) for p in table_patterns)

    while i < n:
        line = lines[i].strip()

        if is_table_line(line):
            # Collect consecutive table lines
            table_lines = []
            while i < n and is_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            # Build table data
            table_data = {}
            for tbl_line in table_lines:
                parts = tbl_line.split(",", 3)
                _, r, c, val = parts
                r, c = int(r), int(c)
                val = val.strip('"').replace('\\\\', '\\').replace('\\"', '"')
                if r not in table_data:
                    table_data[r] = {}
                table_data[r][c] = val

            max_row = max(table_data.keys())
            max_col = max(max(cols.keys()) for cols in table_data.values())

            table = doc.add_table(rows=max_row + 1, cols=max_col + 1)
            table.style = 'Table Grid'

            for r in range(max_row + 1):
                for c in range(max_col + 1):
                    cell_text = table_data.get(r, {}).get(c, "")
                    table.cell(r, c).text = cell_text

            doc.add_paragraph()  # spacing after table

        elif line == "":
            i += 1  # skip empty line
        else:
            # Treat as paragraph
            para_text = line
            p = doc.add_paragraph()

            # Markdown parser for bold/italic
            def parse_markdown_runs(text):
                pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
                last_pos = 0
                for match in pattern.finditer(text):
                    if match.start() > last_pos:
                        yield ('normal', text[last_pos:match.start()])
                    span = match.group()
                    if span.startswith('**'):
                        yield ('bold', span[2:-2])
                    else:
                        yield ('italic', span[1:-1])
                    last_pos = match.end()
                if last_pos < len(text):
                    yield ('normal', text[last_pos:])

            for idx, line_part in enumerate(para_text.split("\n")):
                for run_type, segment in parse_markdown_runs(line_part):
                    run = p.add_run(segment)
                    if run_type == 'bold':
                        run.bold = True
                    elif run_type == 'italic':
                        run.italic = True
                if idx < len(para_text.split("\n")) - 1:
                    p.add_run().add_break()
            i += 1

    # Save document silently
    try:
        doc.save(output_path)
    except Exception as e:
        raise RuntimeError(f"Failed to save Word document: {e}")
