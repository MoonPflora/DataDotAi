import os
import json
import tempfile
from docx import Document

try:
    import win32com.client
except ImportError:
    win32com = None

# --- Helpers ---
def escape_text(text: str) -> str:
    if not text:
        return ""
    return text.replace('\\', '\\\\').replace('"', '\\"').strip()

def create_docx_from_parsed(parsed_text, tmp_docx_path):
    """
    Takes Word-style parsed_text array and writes a DOCX file.
    """
    doc = Document()
    for line in parsed_text:
        line = line.strip()
        if not line:
            continue
        # Tables
        if line.lower().startswith("table"):
            # Parse table row,col,value
            parts = line.split(",", 3)
            if len(parts) == 4:
                _, row_idx, col_idx, cell_text = parts
                row_idx = int(row_idx)
                col_idx = int(col_idx)
                cell_text = cell_text.strip('"')
                # For simplicity, create a new table if first cell
                if doc.tables and len(doc.tables) > 0:
                    table = doc.tables[-1]
                    # expand table if needed
                    while len(table.rows) <= row_idx:
                        table.add_row()
                    while len(table.columns) <= col_idx:
                        for r in table.rows:
                            r.add_cell()
                else:
                    # Create a new table with 1 row/col initially
                    table = doc.add_table(rows=row_idx+1, cols=col_idx+1)
                table.cell(row_idx, col_idx).text = cell_text
        else:
            # Paragraphs
            if line.startswith("paragraph "):
                content = line[len("paragraph "):]
                doc.add_paragraph(content)
            else:
                # fallback
                doc.add_paragraph(line)

    doc.save(tmp_docx_path)


def unparse_pdf(parsed_text, output_pdf_path):
    """
    Takes Word-style parsed_text and writes a proper PDF file.
    """
    if not win32com:
        raise RuntimeError("win32com.client required for PDF conversion on Windows.")

    # Ensure parsed_text is a Python list
    if isinstance(parsed_text, str):
        try:
            parsed_text = json.loads(parsed_text)
        except json.JSONDecodeError:
            parsed_text = [parsed_text]

    # Create temporary DOCX
    tmp_fd, tmp_docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(tmp_fd)
    try:
        create_docx_from_parsed(parsed_text, tmp_docx_path)

        # Convert DOCX → PDF using Word COM
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(tmp_docx_path),
                                      ReadOnly=True,
                                      AddToRecentFiles=False,
                                      Visible=False)
            doc.ExportAsFixedFormat(os.path.abspath(output_pdf_path),
                                    ExportFormat=17)  # 17 = wdExportFormatPDF
            doc.Close(False)
        finally:
            word.Quit()

    finally:
        if os.path.exists(tmp_docx_path):
            os.remove(tmp_docx_path)
