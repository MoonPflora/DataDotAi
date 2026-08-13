import os
import tempfile
import json
from docx import Document
from pdf2docx import Converter

try:
    import win32com.client
except ImportError:
    win32com = None

import fitz  # PyMuPDF
from . import image_parser  # your OCR/image parser


# --- Helpers ---
def escape_text(text: str) -> str:
    """Make text JSON-safe (escape quotes and backslashes)."""
    if not text:
        return ""
    return text.replace('\\', '\\\\').replace('"', '\\"').strip()


# --- PDF as selectable text ---
def parse_pdf_as_word(pdf_path: str) -> list:
    """Parse selectable-text PDF into paragraphs."""
    result = []
    doc = fitz.open(pdf_path)
    for page in doc:
        for line in page.get_text("text").splitlines():
            line = line.strip()
            if line:
                result.append(f"paragraph {escape_text(line)}")
    doc.close()
    return result


# --- Word helpers ---
def parse_paragraph(paragraph) -> str:
    """Parse paragraph text, preserving bold/italic markers."""
    if not paragraph.text.strip():
        return ""
    parts = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        parts.append(text)
    return "".join(parts).strip()


def contains_image(paragraph) -> bool:
    """Check if a paragraph contains an inline image."""
    for run in paragraph.runs:
        for child in run._element:
            if child.tag.endswith('}drawing') or child.tag.endswith('}pict'):
                return True
    return False


def parse_docx_to_array(docx_path: str) -> list:
    """Parse DOCX to JSON-safe array of paragraphs and table cells."""
    doc = Document(docx_path)
    result = []

    # Paragraphs
    for para in doc.paragraphs:
        if contains_image(para):
            result.append("[INSERT IMAGE HERE]")
        text = parse_paragraph(para)
        if text:
            result.append(escape_text(text))

    # Tables
    for t_index, table in enumerate(doc.tables):
        sheet_name = f"Table{t_index}"
        for r_index, row in enumerate(table.rows):
            for c_index, cell in enumerate(row.cells):
                cell_text = escape_text(cell.text)
                has_image = any(contains_image(p) for p in cell.paragraphs)
                if has_image:
                    cell_text = f"{cell_text} [INSERT IMAGE HERE]" if cell_text else "[INSERT IMAGE HERE]"
                result.append(f"{sheet_name},{r_index},{c_index},\"{cell_text}\"")

    return result


def convert_doc_to_docx(doc_path: str) -> str:
    """Convert .doc to .docx using Word COM (Windows only)."""
    if win32com is None:
        raise RuntimeError("win32com.client not installed; cannot convert .doc files on Windows")

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(
            os.path.abspath(doc_path),
            ReadOnly=False,
            ConfirmConversions=False,
            AddToRecentFiles=False,
            Visible=False
        )
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(tmp_fd)
        doc.SaveAs(tmp_path, FileFormat=16)
        doc.Close(False)
    finally:
        word.Quit()
    return tmp_path


# --- PDF detection ---
def pdf_has_text(pdf_path: str, min_lines: int = 2) -> bool:
    """Return True if PDF has enough selectable text lines."""
    try:
        doc = fitz.open(pdf_path)
        line_count = 0
        for page in doc:
            lines = page.get_text("text").splitlines()
            line_count += sum(1 for l in lines if l.strip())
            if line_count >= min_lines:
                doc.close()
                return True
        doc.close()
        return False
    except Exception:
        return False


# --- PDF parsing ---
def parse_pdf(pdf_path: str) -> list:
    """Parse PDF: selectable text → Word parser, scanned → OCR via image_parser."""
    if not os.path.exists(pdf_path):
        return [f"paragraph File not found: {pdf_path}"]

    # Detect selectable text
    if pdf_has_text(pdf_path):
        try:
            return parse_pdf_as_word(pdf_path)
        except Exception as e:
            return [f"paragraph PDF text extraction failed: {str(e)}"]
    else:
        # Scanned PDF → OCR
        try:
            text = image_parser.parse_image(pdf_path, from_pdf=True)
            if not text.strip():
                return ["paragraph "]
            return [f"paragraph {text}"]
        except Exception as e:
            return [f"paragraph OCR failed: {str(e)}"]


# --- Main parser ---
def parse_word(file_path: str) -> list:
    """Parse .doc, .docx, .pdf, .png, .jpg into JSON-safe array of strings."""
    if not os.path.exists(file_path):
        return [f"paragraph File not found: {file_path}"]

    ext = os.path.splitext(file_path)[1].lower()

    if ext in [".pdf"]:
        return parse_pdf(file_path)
    elif ext == ".doc":
        try:
            docx_path = convert_doc_to_docx(file_path)
            return parse_docx_to_array(docx_path)
        except Exception as e:
            return [f"paragraph DOC conversion error: {str(e)}"]
    elif ext == ".docx":
        try:
            return parse_docx_to_array(file_path)
        except Exception as e:
            return [f"paragraph DOCX parsing error: {str(e)}"]
    elif ext in [".png", ".jpg", ".jpeg"]:
        try:
            text = image_parser.parse_image(file_path)
            if not text.strip():
                return ["paragraph "]
            return [f"paragraph {text}"]
        except Exception as e:
            return [f"paragraph OCR failed: {str(e)}"]
    else:
        return [f"paragraph Unsupported file extension: {ext}"]
